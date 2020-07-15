import getopt
import multiprocessing
from multiprocessing import Pool

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import multiprocessing_win
from readword import readword
from screenshot import capture
from word import getdictforword


def main():
   inputfile = input("请输入您要提取的文档路径\n")
   outputfile = input("请输入您要输出的文档路径\n")
   words = readword(inputfile)
   print(words)
   document = Document()
   document.styles["Normal"].font.name=u"Times New Roman"
   document.styles["Normal"].font.element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
   title = document.add_heading('There are words in {}'.format(inputfile.split('.')[0]),level = 0)
   title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
   pool = Pool(processes = 8)
   words_resultlist = []
   for word in words:
      words_resultlist.append(pool.apply_async(getdictforword,args=(word,)))
   pool.close()
   pool.join()
   word_dicts = []
   for word_dict in words_resultlist:
      word_dicts.append(word_dict.get())
   print('开始截屏！')
   i=0
   for word in words:
      capture(word)
      i=i+1
      print('已完成{}%'.format(format((i/len(words))*100,'.2f')))
   i=0
   print('截屏结束！')
   for word in words:
      title_word = document.add_heading(word,level = 1)
      fayin_heading = document.add_heading('发音',level=2)
      fayin_paragraph = document.add_paragraph(word_dicts[i]['发音'])
      mean_heading = document.add_heading('释义',level=2)
      mean_paragraph = document.add_paragraph(word_dicts[i]['释义'])
      try:
         sentence1_heading = document.add_heading('例句1',level=2)
         sentence1_paragraph = document.add_paragraph( word_dicts[i]['例句1'])
         sentence2_heading = document.add_heading('例句2',level=2)
         sentence2_paragraph = document.add_paragraph( word_dicts[i]['例句2'])
      except:
         pass
      document.add_picture('images\\{}.png'.format(word),width = Inches(6.0))
      document.add_page_break()
      i+=1
   document.save(outputfile)
   print('finished')

if __name__ == "__main__":
   multiprocessing.freeze_support()
   main()
